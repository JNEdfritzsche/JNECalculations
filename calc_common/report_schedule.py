from __future__ import annotations

import math
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Callable

from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
import pandas as pd
import streamlit as st
from streamlit.components.v1 import html as components_html

from calc_common.formatting import Quantity, fmt
from calc_common.report_helper import (
    add_word_table,
    init_word_doc,
    render_export_buttons,
    report_identity,
    save_word_doc,
    wb_to_bytes,
)

ValueGetter = Callable[[dict[str, Any]], Any]
WordRefBuilder = Callable[[Any, list[dict[str, Any]]], None]

INPUT_HEADER_FILL = PatternFill("solid", fgColor="EDEDED")
RESULT_HEADER_FILL = PatternFill("solid", fgColor="D6E4F0")
RESULT_FONT_COLORS = {"green": "1B7F3B", "blue": "1F4E79"}

RESULTS_SHEET = "Results"
INPUTS_SHEET = "Inputs"

BLANK = (None, "", "—")


@dataclass
class Group:
    """A set of columns the Word report prints in one cell. Excel keeps them apart."""
    label: str
    sep: str = " / "


@dataclass
class Column:
    header: str
    get: ValueGetter
    color: str | None = None  # export font colour: green = the answer, blue = the key selection
    result: bool = False  # value the calculator derived, as opposed to one the designer entered
    group: str | None = None  # key into ReportSpec.groups — Word prints the group in one cell
    always: bool = False  # keep the column even when every row shares one value


@dataclass
class ReportSpec:
    code: str # jurisdiction (eg. nec, oesc)
    calculator: str # package directory name (eg. voltage_drop)
    report_title: str # eg. Voltage Drop — Calculation Results
    sheet_name: str # name for the excel sheet
    cols: list[Column | tuple[str, ValueGetter]]  # Column, or legacy (header, getter) tuple
    code_reference: Callable[[list[dict[str, Any]]], str]  # "per NEC …" line from queued results
    notes: Callable[[list[dict[str, Any]]], list[str]]  # note lines under the table
    tag: str = "Tag"
    word_reference: WordRefBuilder | None = None  # appends equations/assumptions to the Word report
    groups: dict[str, Group] = field(default_factory=dict)  # Word-only cell merging
    input_prefixes: tuple[str, ...] | None = None  # widget key prefixes to snapshot; () disables Edit
    prefix: str = field(init=False)  # namespaces every session_state and widget key

    def __post_init__(self):
        self.prefix = f"{self.code}_{self.calculator}_schedule"
        if self.input_prefixes is None:
            self.input_prefixes = (f"{self.code}_{self.calculator}_",)


# ============================================================
# st.session_state helpers
# ============================================================
def _rows_key(spec: ReportSpec) -> str:
    return f"{spec.prefix}_rows"

def _table_key(spec: ReportSpec) -> str:
    return f"{spec.prefix}_table"

def _remove_row_by_index(spec: ReportSpec, index: int):
    rows = get_rows(spec)
    if 0 <= index < len(rows):
        rows.pop(index)

def _remove_row_by_tag(spec: ReportSpec, tag: str):
    rows = get_rows(spec)
    for i, row in enumerate(rows):
        if row["tag"] == tag:
            rows.pop(i)
            break

def get_rows(spec: ReportSpec) -> list[dict[str, Any]]:
    key = _rows_key(spec)
    if key not in st.session_state:
        st.session_state[key] = []
    return st.session_state[key]

# Streamlit refuses a session_state write for these, and raises when the widget is
# next created rather than at the write, so they have to be kept out of the snapshot.
UNSETTABLE_VALUE_TYPES = {
    "trigger_value", "string_trigger_value", "json_trigger_value",
    "chat_input_value", "file_uploader_state_value", "bytes_value",
}

def _unsettable(keys: list[str]) -> set[str]:
    # The widget type isn't derivable from the key, so it is read off Streamlit's
    # widget metadata. Only valid while the widgets are live — metadata for another
    # page has already been discarded.
    try:
        from streamlit.runtime.state import get_session_state

        state = get_session_state()._state
        metadata = state._new_widget_state.widget_metadata
        mapper = state._key_id_mapper
        return {
            key for key in keys
            if (m := metadata.get(mapper.get_id_from_key(key))) is not None
            and m.value_type in UNSETTABLE_VALUE_TYPES
        }
    except Exception:
        # internals moved: drop the shapes a trigger or uploader would hold
        return {k for k in keys if st.session_state.get(k) is False
                or st.session_state.get(k) is None}

def _snapshot_inputs(spec: ReportSpec) -> dict[str, Any]:
    # spec.prefix itself starts with the widget prefix, so without excluding it the
    # row list would be snapshotted into one of its own rows.
    keys = [
        key for key in st.session_state
        if any(key.startswith(p) for p in spec.input_prefixes)
        and not key.startswith(spec.prefix)
    ]
    skip = _unsettable(keys)
    return {key: st.session_state[key] for key in keys if key not in skip}

def add_row(spec: ReportSpec, tag: str, result: dict[str, Any]):
    get_rows(spec).append({"tag": tag, "result": result, "inputs": _snapshot_inputs(spec)})

def find_row(spec: ReportSpec, tag: str) -> dict[str, Any] | None:
    return next((r for r in get_rows(spec) if r["tag"] == tag), None)

def replace_row(spec: ReportSpec, tag: str, result: dict[str, Any]) -> None:
    for row in get_rows(spec):
        if row["tag"] == tag:
            row["result"] = result
            row["inputs"] = _snapshot_inputs(spec)
            return
    add_row(spec, tag, result)

def remove_row(spec: ReportSpec, identifier: str | int):
    if isinstance(identifier, int):
        _remove_row_by_index(spec, identifier)
    else:
        _remove_row_by_tag(spec, identifier)

def clear(spec: ReportSpec) -> None:
    st.session_state[_rows_key(spec)] = []

def _keep_key(spec: ReportSpec) -> str:
    return f"{spec.prefix}_keep"

def apply_restore(spec: ReportSpec) -> None:
    # Streamlit drops a widget's state on any run where the widget isn't rendered, so
    # navigating to another calculator and back would reset every input to its default.
    # Only absent keys are filled, so a live value is never clobbered.
    for key, value in st.session_state.get(_keep_key(spec), {}).items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Widget state can only be written before the widget is instantiated, so the Edit
    # button only queues the snapshot and this runs it at the top of the next pass.
    pending = st.session_state.pop(f"{spec.prefix}_restore", None)
    if not pending:
        return
    for key, value in pending["inputs"].items():
        st.session_state[key] = value
    st.session_state[f"{spec.prefix}_tag"] = pending["tag"]


# ============================================================
# row builder helpers
# ============================================================

def _as_columns(spec: ReportSpec) -> list[Column]:
    """Normalize spec.cols (Column objects and/or legacy (header, getter) tuples)."""
    return [c if isinstance(c, Column) else Column(c[0], c[1]) for c in spec.cols]

def _safe(get: ValueGetter, result: dict[str, Any]) -> Any:
    try:
        val = get(result)
    except Exception:
        return "—"
    return val if val not in (None, "") else "—"

def _grouped_columns(spec: ReportSpec) -> list[Column]:
    # A group header spans one contiguous run, so interleaved results and inputs
    # would draw as alternating Results/Inputs bands rather than two.
    columns = _as_columns(spec)
    return [c for c in columns if c.result] + [c for c in columns if not c.result]

def _headers(spec: ReportSpec) -> list[str]:
    return [spec.tag] + [c.header for c in _grouped_columns(spec)]

def _grouped_headers(spec: ReportSpec) -> pd.MultiIndex:
    # Streamlit draws the upper level as a spanning group header. An empty upper
    # level is dropped rather than grouped under, which is what the tag wants.
    return pd.MultiIndex.from_tuples(
        [("", spec.tag)]
        + [(RESULTS_SHEET if c.result else INPUTS_SHEET, c.header)
           for c in _grouped_columns(spec)]
    )

def _rows_for(spec: ReportSpec, rows: list[dict[str, Any]], columns: list[Column]) -> list[dict[str, Any]]:
    table_rows: list[dict[str, Any]] = []
    for row in rows:
        result = row.get("result", {})
        data_row: dict[str, Any] = {spec.tag: row.get("tag", "")}
        for c in columns:
            data_row[c.header] = _safe(c.get, result)
        table_rows.append(data_row)
    return table_rows

def _build_table_rows(spec: ReportSpec, rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return _rows_for(spec, rows, _grouped_columns(spec))

def _split_columns(spec: ReportSpec) -> list[tuple[str, list[Column]]]:
    """Results first — the answers are what the schedule is read for."""
    columns = _as_columns(spec)
    return [
        (RESULTS_SHEET, [c for c in columns if c.result]),
        (INPUTS_SHEET, [c for c in columns if not c.result]),
    ]

def _partition(columns: list[Column], table_rows: list[dict[str, Any]]) -> tuple[list[Column], list[tuple[str, Any]]]:
    """Columns that vary, and the (header, value) pairs that are the same on every row.

    A column carrying one repeated value says nothing a single line below the table
    cannot. Headline columns are exempt — burying the answer in a footnote is wrong
    even when every row agrees. With fewer than two rows nothing is meaningfully
    constant, so everything is kept.
    """
    if len(table_rows) < 2:
        return columns, []

    kept: list[Column] = []
    constant: list[tuple[str, Any]] = []
    for column in columns:
        values = [row.get(column.header) for row in table_rows]
        if all(v in BLANK for v in values):
            # An empty column states nothing, so there is nothing to hoist either.
            continue
        if column.always or column.color:
            kept.append(column)
        elif len({str(v) for v in values}) == 1:
            constant.append((column.header, values[0]))
        else:
            kept.append(column)
    return kept, constant

def _joined(members: list[Column], sep: str) -> ValueGetter:
    def get(result: dict[str, Any]) -> str:
        parts = [str(_safe(m.get, result)) for m in members]
        # "Not required → Not required" says it twice.
        return parts[0] if len(set(parts)) == 1 else sep.join(parts)
    return get

def _grouped(columns: list[Column], spec: ReportSpec) -> list[Column]:
    """Word only — collapse each declared group into a single column."""
    members: dict[str, list[Column]] = {}
    for column in columns:
        if column.group:
            members.setdefault(column.group, []).append(column)

    out: list[Column] = []
    done: set[str] = set()
    for column in columns:
        key = column.group
        if not key or key not in spec.groups:
            out.append(column)
            continue
        if key in done:
            continue
        done.add(key)
        group_columns = members[key]
        # A group whose other members were hoisted out is just that one column, and
        # labelling a lone trade size "Conduit" would misread.
        if len(group_columns) == 1:
            out.append(group_columns[0])
            continue
        group = spec.groups[key]
        out.append(Column(
            group.label,
            _joined(group_columns, group.sep),
            color=next((c.color for c in group_columns if c.color), None),
            result=any(c.result for c in group_columns),
        ))
    return out

CONSTANTS_HEADING = "Same for all rows"


def _write_constants_word(doc, constant: list[tuple[str, Any]]) -> None:
    # Bold the input name so it reads apart from the value it carries.
    para = doc.add_paragraph()
    for index, (header, value) in enumerate(constant):
        prefix = f"{CONSTANTS_HEADING}:   " if index == 0 else "   ·   "
        for text, bold in ((prefix, index == 0), (f"{header}: ", True), (str(value), False)):
            run = para.add_run(text)
            run.bold = bold
            run.italic = True
            run.font.size = Pt(8)


# ============================================================
# Excel cell helpers
# ============================================================

def _numbers(values: list[Any]) -> list[float]:
    out = []
    for value in values:
        if isinstance(value, Quantity):
            out.append(value.number)
        elif isinstance(value, (int, float)) and not isinstance(value, bool):
            out.append(float(value))
    return out

def _column_unit(values: list[Any]) -> str | None:
    """The one unit a column carries, or None where its rows disagree."""
    units = {v.unit for v in values if isinstance(v, Quantity) and v.unit}
    return units.pop() if len(units) == 1 else None

def _excel_header(header: str, unit: str | None, numeric: bool) -> str:
    if not numeric or not unit or f"({unit})" in header:
        return header
    return f"{header} ({unit})"

def _excel_value(value: Any, numeric: bool, has_numbers: bool) -> Any:
    if isinstance(value, Quantity):
        return value.number if numeric else str(value)
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, (int, float)):
        return value
    if value in (None, ""):
        return None
    # A blank keeps formulas clean where the column carries numbers; in a text
    # column the dash is the value, so blanking it just looks like missing data.
    if value == "—":
        return None if has_numbers else "—"
    return str(value)

WHOLE_FORMAT = "#,##0"


def _number_format(values: list[Any]) -> str | None:
    """One format per column, sized to its smallest value so nothing renders as an exponent."""
    magnitudes = [abs(n) for n in _numbers(values) if n]
    if not magnitudes:
        return None
    decimals = min(8, max(0, 3 - math.floor(math.log10(min(magnitudes)))))
    return WHOLE_FORMAT + ("." + "#" * decimals if decimals else "")


def _cell_format(value: float, column_format: str) -> str:
    """Excel prints the decimal separator whenever the format carries one, so 30 under
    #,##0.### reads '30.'. A whole value takes the plain integer format instead."""
    return WHOLE_FORMAT if float(value).is_integer() else column_format

MIN_COL_WIDTH = 9
MAX_COL_WIDTH = 34
MAX_HEADER_LINES = 4


def _display_len(value: Any) -> int:
    # A float's repr runs far longer than the number format shows it as.
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        return len(str(value or ""))
    return len(str(fmt(value)))


def _column_width(header: str, cells: list[Any]) -> int:
    """Size to the data, not the header — headers wrap, so a long one like
    'Sec. breaker selected (A)' should not stretch a column of 3-digit values."""
    widest_data = max((_display_len(c) for c in cells), default=0)
    longest_word = min(max((len(w) for w in header.split()), default=0), 14)
    return max(MIN_COL_WIDTH, min(MAX_COL_WIDTH, max(widest_data, longest_word) + 2))


def _write_schedule_table(ws, header_row: int, tag: str, cols: list[Column],
                          table_rows: list[dict[str, Any]]) -> None:
    headers = [tag] + [c.header for c in cols]
    columns: list[Column | None] = [None, *cols]  # the tag column has no Column
    header_lines = 1

    for col_idx, (header, column) in enumerate(zip(headers, columns), start=1):
        values = [row.get(header) for row in table_rows]
        unit = _column_unit(values)
        # Mixed units in one column can't be hoisted into the header, so that column stays text.
        numeric = unit is not None or not any(isinstance(v, Quantity) and v.unit for v in values)
        has_numbers = bool(_numbers(values))
        is_result = bool(column and column.result)

        label = _excel_header(header, unit, numeric)
        cell = ws.cell(row=header_row, column=col_idx, value=label)
        cell.font = Font(bold=True)
        cell.fill = RESULT_HEADER_FILL if is_result else INPUT_HEADER_FILL
        cell.alignment = Alignment(wrap_text=True, vertical="bottom")

        number_format = _number_format(values) if numeric else None
        color = RESULT_FONT_COLORS.get(column.color) if column else None
        font = Font(bold=is_result, color=color) if (is_result or color) else None

        written = []
        for offset, value in enumerate(values, start=1):
            data = ws.cell(row=header_row + offset, column=col_idx,
                           value=_excel_value(value, numeric, has_numbers))
            written.append(data.value)
            if font is not None:
                data.font = font
            if number_format and isinstance(data.value, (int, float)):
                data.number_format = _cell_format(data.value, number_format)

        width = _column_width(label, written)
        ws.column_dimensions[get_column_letter(col_idx)].width = width
        header_lines = max(header_lines, min(MAX_HEADER_LINES, -(-len(label) // max(1, width - 1))))

    ws.row_dimensions[header_row].height = 14 * header_lines
    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)
    if table_rows:
        last = get_column_letter(len(headers))
        ws.auto_filter.ref = f"A{header_row}:{last}{header_row + len(table_rows)}"

def _fit_column(ws, col_idx: int, widest: int) -> None:
    # Never shrink: the table above sized these same columns for its own data.
    letter = get_column_letter(col_idx)
    current = ws.column_dimensions[letter].width or 0
    ws.column_dimensions[letter].width = max(current, min(MAX_COL_WIDTH, widest + 2))


def _write_constants_excel(ws, row: int, constant: list[tuple[str, Any]]) -> int:
    # Column A is always the input, column B its value. Returns the row below the block.
    heading = ws.cell(row=row, column=1, value=CONSTANTS_HEADING)
    heading.font = Font(bold=True)
    row += 1

    labels, values = [], []
    for header, value in constant:
        label = ws.cell(row=row, column=1, value=header)
        label.font = Font(bold=True, size=9)
        label.fill = INPUT_HEADER_FILL
        cell = ws.cell(row=row, column=2, value=_excel_value(value, False, False))
        cell.font = Font(size=9)
        labels.append(len(header))
        values.append(_display_len(cell.value))
        row += 1

    # The heading is left unfitted; its row leaves column B empty so it can overflow.
    _fit_column(ws, 1, max(labels))
    _fit_column(ws, 2, max(values))
    return row


def _set_landscape(doc) -> None:
    """Widen to landscape so the multi-column schedule table isn't cramped."""
    try:
        section = doc.sections[0]
        width, height = section.page_width, section.page_height
        if width is None or height is None:
            return
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = max(width, height), min(width, height)
    except Exception:
        pass

def build_schedule_word(spec: ReportSpec, rows: list[dict[str, Any]]):
    results = [row.get("result", {}) for row in rows]
    doc = init_word_doc(spec.report_title)
    _set_landscape(doc)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Generated {datetime.now().strftime('%B %d, %Y · %I:%M %p')}   |   {spec.code_reference(results)}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)

    for title, columns in _split_columns(spec):
        kept, constant = _partition(columns, _rows_for(spec, rows, columns))
        if not kept and not constant:
            continue

        heading = doc.add_paragraph().add_run(title)
        heading.bold = True
        heading.font.size = Pt(10)

        if kept:
            display = _grouped(kept, spec)
            add_word_table(
                doc,
                [spec.tag] + [c.header for c in display],
                _rows_for(spec, rows, display),
                font_size_header=8,
                font_size=8,
                bold_columns={c.header for c in display if c.result},
            )

        if constant:
            _write_constants_word(doc, constant)

    # Calculator-supplied reference content (equations, etc.).
    if spec.word_reference:
        spec.word_reference(doc, results)

    for note in spec.notes(results):
        para = doc.add_paragraph()
        note_run = para.add_run(f"Note: {note}")
        note_run.italic = True
        note_run.font.size = Pt(8)

    return save_word_doc(doc)

def _write_identity(ws) -> int:
    row = 1
    for label, value in zip(("Project No.", "Designer"), report_identity()):
        ws.cell(row=row, column=1, value=label).font = Font(bold=True)
        ws.cell(row=row, column=2, value=value)
        row += 1
    return row + 1

def build_schedule_excel(spec: ReportSpec, rows: list[dict[str, Any]]):
    results = [row.get("result", {}) for row in rows]

    wb = Workbook()
    sheets = _split_columns(spec)

    for index, (title, columns) in enumerate(sheets):
        ws = wb.active if index == 0 else wb.create_sheet()
        ws.title = title

        row = _write_identity(ws)
        kept, constant = _partition(columns, _rows_for(spec, rows, columns))
        table_rows = _rows_for(spec, rows, kept)

        _write_schedule_table(ws, row, spec.tag, kept, table_rows)
        row += len(table_rows) + 2

        if constant:
            row = _write_constants_excel(ws, row, constant) + 1

        if index:
            continue

        meta = ws.cell(row=row, column=1)
        meta.value = (
            f"Generated {datetime.now().strftime('%B %d, %Y · %I:%M %p')}"
            f"   |   {spec.code_reference(results)}"
        )
        meta.font = Font(bold=True)
        row += 2

        for note in spec.notes(results):
            cell = ws.cell(row=row, column=1, value=f"Note: {note}")
            cell.font = Font(italic=True, size=9)
            row += 1

    return wb_to_bytes(wb)
    
# ============================================================
# UI
# ============================================================

def render_schedule_commit(spec: ReportSpec, result: dict[str, Any] | None, can_add: bool) -> None:
    rows = get_rows(spec)
    placeholder = f"Calc {len(rows) + 1}"
    tag_key = f"{spec.prefix}_tag"

    # clear_on_submit would wipe a tag written by Edit, so the reset is queued instead
    # and applied here, before the widget is instantiated
    if st.session_state.pop(f"{spec.prefix}_clear_tag", False):
        st.session_state[tag_key] = ""

    with st.form(f"{spec.prefix}_add_form", border=False):
        tag = st.text_input(spec.tag, key=tag_key, placeholder=placeholder)
        submitted = st.form_submit_button(
            "Add to schedule",
            width="stretch",
            disabled=not can_add,
            type="primary"
        )

    if not can_add:
        st.caption("Enter valid inputs to add this calculation to the schedule.")
    elif rows:
        st.caption(f"{len(rows)} in schedule")

    pending_key = f"{spec.prefix}_pending"

    if submitted and result is not None:
        name = (tag or "").strip() or placeholder
        if find_row(spec, name):
            st.session_state[pending_key] = {"tag": name, "result": result}
        else:
            add_row(spec, name, result)
            st.session_state[f"{spec.prefix}_clear_tag"] = True
            st.rerun()

    pending = st.session_state.get(pending_key)
    if pending:
        st.warning(f"**{pending['tag']}** is already in the schedule.")
        keep, cancel = st.columns(2)
        if keep.button("Overwrite it", key=f"{spec.prefix}_overwrite",
                       type="primary", width="stretch"):
            replace_row(spec, pending["tag"], pending["result"])
            del st.session_state[pending_key]
            st.session_state[f"{spec.prefix}_clear_tag"] = True
            st.rerun()
        if cancel.button("Cancel", key=f"{spec.prefix}_cancel_overwrite", width="stretch"):
            del st.session_state[pending_key]
            st.rerun()


def render_schedule_table(spec: ReportSpec) -> None:
    rows = get_rows(spec)

    # Every input has rendered by this point, so this is where the mirror that survives
    # navigation is taken. Above the empty-schedule return, which fires on first visit.
    st.session_state[_keep_key(spec)] = _snapshot_inputs(spec)

    st.markdown("### Report schedule")

    if not rows:
        st.caption("Nothing queued yet. Set up a circuit above, then use **Add to schedule**.")
        return
    c1, _, c2 = st.columns(3)
    c1.caption(f"{len(rows)} calculation{'s' if len(rows) != 1 else ''} queued — select rows to remove them.")
    
    can_edit = bool(spec.input_prefixes)
    controls = c2.columns(3 if can_edit else 2)
    edit_col = controls[0] if can_edit else None
    remove_col, clear_col = controls[-2], controls[-1]
    
    table = pd.DataFrame(_build_table_rows(spec, rows), columns=_headers(spec))
    table.columns = _grouped_headers(spec)

    state = st.dataframe(
        table,
        width="stretch",
        on_select="rerun",
        selection_mode="multi-row",
        key=_table_key(spec),
    )
    selected = list(getattr(state, "selection", {}).get("rows", []))

    hoisted = sum(
        len(_partition(columns, _rows_for(spec, rows, columns))[1])
        for _title, columns in _split_columns(spec)
    )
    note = "Results and inputs export as separate tables."
    if hoisted:
        note += f" {hoisted} column{'s' if hoisted != 1 else ''} identical on every row move to a note."
    st.caption(note)


    editable = rows[int(selected[0])] if len(selected) == 1 else None
    if edit_col is not None and edit_col.button(
        "Edit selected",
        key=f"{spec.prefix}_edit_selected",
        disabled=not (editable and editable.get("inputs")),
        width="stretch",
    ):
        st.session_state[f"{spec.prefix}_restore"] = {
            "inputs": editable["inputs"],
            "tag": editable["tag"],
        }
        st.session_state[f"{spec.prefix}_scroll"] = True
        st.rerun()

    if remove_col.button(
        f"Remove selected ({len(selected)})",
        key=f"{spec.prefix}_remove_selected",
        disabled=not selected,
        width="stretch",
    ):
        for index in sorted(selected, reverse=True):
            remove_row(spec, int(index))
        st.rerun()

    if clear_col.button("Clear all", key=f"{spec.prefix}_clear_all", width="stretch"):
        clear(spec)
        st.rerun()
    
    

    render_export_buttons(
        prefix=spec.prefix,
        docx_file=f"{spec.prefix}.docx",
        xlsx_file=f"{spec.prefix}.xlsx",
        can_export=True,
        word_builder=lambda: build_schedule_word(spec, get_rows(spec)),
        excel_builder=lambda: build_schedule_excel(spec, get_rows(spec)),
    )


def render_schedule_ui(spec: ReportSpec, result: dict[str, Any] | None, can_add: bool):
    st.markdown("### Add to report")
    render_schedule_commit(spec, result, can_add)
    st.divider()
    render_schedule_table(spec)
from __future__ import annotations

import io
from datetime import datetime
from html import escape
from pathlib import Path

import streamlit as st
from calc_common.formatting import fmt
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from openpyxl import Workbook


# ============================================================
# Value formatting
# ============================================================
def cell_text(val):
    if val is None or val == "":
        return "—"
    if isinstance(val, bool):
        return str(val)
    if isinstance(val, (int, float)):
        return str(fmt(val))
    return str(val)


def get_first(data, *keys, default=None):
    for key in keys:
        if key in data and data.get(key) is not None:
            return data.get(key)
    return default


def yes_no(value):
    return "Yes" if bool(value) else "No"


# ============================================================
# Word equation helpers
# ============================================================
def omml_r(text: str) -> str:
    return f"<m:r><m:t>{escape(str(text))}</m:t></m:r>"


def omml_sub(base: str, sub: str) -> str:
    return f"<m:sSub><m:e>{omml_r(base)}</m:e><m:sub>{omml_r(sub)}</m:sub></m:sSub>"


def omml_frac(num_inner: str, den_inner: str) -> str:
    return f"<m:f><m:num>{num_inner}</m:num><m:den>{den_inner}</m:den></m:f>"


def omml_sqrt(inner: str) -> str:
    return f'<m:rad><m:degHide m:val="1"/><m:e>{inner}</m:e></m:rad>'


def add_omml_equation_to_paragraph(p, omml_inner: str) -> None:
    xml = f'<m:oMath {nsdecls("m")}>{omml_inner}</m:oMath>'
    p._p.append(parse_xml(xml))


def add_word_equation(doc, label: str, omml_inner: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    add_omml_equation_to_paragraph(p, omml_inner)


# ============================================================
# Word document helpers
# ============================================================
def append_to_value_line(cell, value: str, paragraph_index: int = 1):
    while len(cell.paragraphs) <= paragraph_index:
        cell.add_paragraph("")
    cell.paragraphs[paragraph_index].add_run(value)


def _delete_paragraph(p):
    p._element.getparent().remove(p._element)
    p._p = p._element = None


def remove_leading_blank_paragraphs(doc: Document):
    while doc.paragraphs and doc.paragraphs[0].text.strip() == "":
        _delete_paragraph(doc.paragraphs[0])


def set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    table._tbl.tblPr.append(borders)


def report_identity() -> tuple[str, str]:
    # (project number, designer name) as entered in the sidebar 
    return (
        st.session_state.get("project_number", ""),
        st.session_state.get("designer_name", ""),
    )


def add_identity_line(doc: Document) -> None:
    project_number, designer_name = report_identity()
    parts = [
        text
        for text in (
            f"Project No.: {project_number}" if project_number else "",
            f"Designer: {designer_name}" if designer_name else "",
            datetime.now().strftime("%B %d, %Y"),
        )
        if text
    ]
    run = doc.add_paragraph().add_run("   |   ".join(parts))
    run.italic = True
    run.font.size = Pt(9)


def fill_doc_header(doc: Document, title: str) -> None:
    project_number, designer_name = report_identity()
    values = {
        (0, 3): project_number,
        (0, 4): "#",
        (2, 3): designer_name,
        (2, 4): datetime.now().strftime("%m/%d/%Y"),
        (3, 2): title,
        (3, 3): "",
        (3, 4): "",
    }
    try:
        hdr_table = doc.sections[0].header.tables[0]
        for (row, col), value in values.items():
            append_to_value_line(hdr_table.cell(row, col), value)
    except Exception:
        doc.add_heading(title, level=0)


def init_word_doc(title: str, template_path: str = "content/files/Template.docx") -> Document:
    path = Path(template_path)
    if not path.exists():
        doc = Document()
        doc.add_heading(title, level=0)
        add_identity_line(doc)
        return doc

    doc = Document(str(path))
    remove_leading_blank_paragraphs(doc)
    fill_doc_header(doc, title)
    return doc


def save_word_doc(doc: Document) -> bytes:
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _set_cell_text(cell, value, font_size, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(cell_text(value))
    run.bold = bold
    run.font.size = Pt(font_size)


def _repeat_header_row(row) -> None:
    """Reprint the header on every page — a schedule easily runs past one."""
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(header)


def _tighten_cells(table, margin_twips: int = 36) -> None:
    """Shrink the default cell padding so a wide schedule fits the page."""
    margins = OxmlElement("w:tblCellMar")
    for name in ("left", "right"):
        element = OxmlElement(f"w:{name}")
        element.set(qn("w:w"), str(margin_twips))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    table._tbl.tblPr.append(margins)


def add_word_table(doc, headers, rows, font_size_header=9, font_size=9, bold_columns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    bold_columns = bold_columns or set()

    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, font_size_header, bold=True)
    _repeat_header_row(table.rows[0])

    for row_data in rows:
        cells = table.add_row().cells
        values = [
            row_data.get(header) if isinstance(row_data, dict) else row_data[i] if i < len(row_data) else None
            for i, header in enumerate(headers)
        ]
        for cell, header, value in zip(cells, headers, values):
            _set_cell_text(cell, value, font_size, bold=header in bold_columns)

    set_table_borders(table)
    _tighten_cells(table)
    return table


# ============================================================
# Excel helpers
# ============================================================
def wb_to_bytes(wb: Workbook) -> bytes:
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


# ============================================================
# Export button pair
# ============================================================
def render_export_buttons(prefix, docx_file, xlsx_file, can_export, word_builder, excel_builder):
    # Built eagerly, during the script run. st.download_button also accepts a callable
    # for deferred building, but Streamlit invokes it without a script context, so
    # anything reading st.session_state (queued rows, the letterhead's project number
    # and designer name) comes back empty and the report exports blank.
    exports = [
        (
            "docx",
            "Word",
            docx_file,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            word_builder,
        ),
        (
            "xlsx",
            "Excel",
            xlsx_file,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            excel_builder,
        ),
    ]

    for col, (ext, label, file_name, mime, builder) in zip(st.columns([1, 1], gap="large"), exports):
        data = b""
        failed = False

        if can_export:
            try:
                data = builder()
            except Exception as exc:
                col.error(f"Could not build the {label} report: {exc}")
                failed = True

        col.download_button(
            f"⬇️ Download {label} report (.{ext})",
            data=data,
            file_name=file_name,
            mime=mime,
            disabled=(not can_export) or failed,
            key=f"{prefix}_download_{ext}",
            width="stretch",
        )


__all__ = [
    # Value formatting
    "get_first",
    "yes_no",

    # Word equation helpers
    "omml_r",
    "omml_sub",
    "omml_frac",
    "omml_sqrt",
    "add_omml_equation_to_paragraph",
    "add_word_equation",

    # Word document helpers
    "report_identity",
    "add_identity_line",
    "init_word_doc",
    "save_word_doc",
    "add_word_table",

    # Excel helpers
    "wb_to_bytes",

    # Export button pair
    "render_export_buttons",
]

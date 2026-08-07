from __future__ import annotations

from html import escape
from typing import Any

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl.styles import Font

from nec_calc.common.report_helper import (
    add_bullets,
    add_kv_section_to_word,
    add_source_table_to_word,
    autosize_cols,
    build_nec_table_row_source,
    cell_text,
    get_first,
    init_excel_report,
    init_word_doc,
    notes_to_pairs,
    render_export_buttons,
    save_word_doc,
    wb_to_bytes,
    write_columnar_table,
    write_kv_sections_to_excel,
    write_source_table_to_excel,
    yes_no,
)


REPORT_TITLE = "NEC Conductor Ampacity & Derating Calculation Report"

METHOD_LABELS = {
    "verify": "Ampacity Verification (Derating & Terminal Check)",
    "size": "Minimum Conductor Sizing",
}

MATERIAL_LABELS = {
    "cu": "Copper",
    "al": "Aluminum or Copper-Clad Aluminum",
}


def _method_label(calc_mode: str) -> str:
    return METHOD_LABELS.get(calc_mode, str(calc_mode))


def _omml_r(text: str) -> str:
    return f"<m:r><m:t>{escape(str(text))}</m:t></m:r>"


def _omml_sub(base: str, sub: str) -> str:
    return (
        "<m:sSub>"
        f"<m:e>{_omml_r(base)}</m:e>"
        f"<m:sub>{_omml_r(sub)}</m:sub>"
        "</m:sSub>"
    )


def _add_omml_equation_to_paragraph(p, omml_inner: str) -> None:
    xml = f'<m:oMath {nsdecls("m")}>{omml_inner}</m:oMath>'
    p._p.append(parse_xml(xml))


def _add_word_equation(doc, label: str, omml_inner: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    _add_omml_equation_to_paragraph(p, omml_inner)


def add_equations(doc, calc_mode: str) -> None:
    doc.add_heading("Equations Used", level=1)

    eq1 = (
        _omml_sub("I", "derated")
        + _omml_r(" = ")
        + _omml_sub("I", "table")
        + _omml_r(" × ")
        + _omml_sub("CF", "temp")
        + _omml_r(" × ")
        + _omml_sub("AF", "cond")
    )

    eq2 = (
        _omml_sub("I", "allowable")
        + _omml_r(" = min(")
        + _omml_sub("I", "derated")
        + _omml_r(", ")
        + _omml_sub("I", "terminal")
        + _omml_r(")")
    )

    _add_word_equation(doc, "Derated Ampacity (NEC 310.15)", eq1)
    _add_word_equation(doc, "Final Allowable Ampacity (NEC 110.14(C))", eq2)


def _build_equations_for_excel(calc_mode: str) -> list[tuple[str, str]]:
    return [
        ("Derated Ampacity (NEC 310.15)", "I_derated = I_table × CF_temp × AF_cond"),
        ("Final Allowable Ampacity (NEC 110.14(C))", "I_allowable = min(I_derated, I_terminal)"),
    ]


def _build_result_pairs(result: dict[str, Any]) -> list[tuple[str, Any]]:
    pairs: list[tuple[str, Any]] = [
        ("Selected Conductor Size", get_first(result, "selected_size_display", "selected_size", "conductor_size_display", "conductor_size")),
    ]
    load_current = get_first(result, "load_current")
    if load_current is not None:
        min_size = get_first(result, "min_recommended_size_display", default="—")
        if min_size and min_size != "—":
            pairs.append(("Minimum Required Size", min_size))

    pairs.extend([
        ("Final Allowable Ampacity (A)", get_first(result, "calculated_value")),
        ("Derated Ampacity per Conductor (A)", get_first(result, "derated_ampacity")),
        ("Table 310.16 Base Ampacity (A)", get_first(result, "table_ampacity")),
        ("Ambient Correction Factor (CF)", get_first(result, "ambient_correction")),
        ("Adjustment Factor (AF)", get_first(result, "conductor_adjustment")),
    ])

    term_limit = get_first(result, "terminal_limit_ampacity")
    if term_limit is not None and term_limit != "—":
        pairs.append(("Terminal Limit Ampacity (A)", term_limit))

    if load_current is not None:
        pairs.append(("Adequate for Load?", yes_no(get_first(result, "is_adequate"))))

    return pairs


def _build_input_pairs(
    result: dict[str, Any],
    calc_mode: str,
) -> list[tuple[str, Any]]:
    mat = get_first(result, "material")
    pairs: list[tuple[str, Any]] = [
        ("Calculation Mode", "Unified Ampacity & Verification Mode"),
    ]
    load_current = get_first(result, "load_current")
    if load_current is not None:
        pairs.append(("Design Load Current (A)", load_current))

    pairs.append(("Conductor Material", MATERIAL_LABELS.get(mat, str(mat))))

    wire_type = get_first(result, "wire_type")
    if wire_type and wire_type != "Not specified":
        pairs.append(("Wire / Cable Type", wire_type))

    n_parallel = get_first(result, "n_parallel", default=1)
    if n_parallel and n_parallel > 1:
        pairs.append(("Parallel Runs per Phase (N)", n_parallel))

    pairs.append(("Conductor Insulation Rating (°C)", get_first(result, "temp_rating")))

    terminal_rating = get_first(result, "terminal_temp_rating")
    if terminal_rating is not None and terminal_rating != "—" and terminal_rating != "":
        pairs.append(("Terminal Temperature Rating (°C)", terminal_rating))

    pairs.append(("Ambient Temperature Range (°C)", get_first(result, "ambient_temp_c")))
    pairs.append(("Number of Current-Carrying Conductors", get_first(result, "number_of_conductors")))
    return pairs


def _build_notes(calc_mode: str) -> list[str]:
    return [
        "Conductor ampacity calculations are performed in accordance with NEC Articles 310 and 110.14(C).",
        "Base ampacities are retrieved from NEC Table 310.16 (based on 30°C ambient and not more than 3 current-carrying conductors).",
        "Ambient temperature correction factors are determined from Table 310.15(B)(1)(1) or (2) based on conductor insulation rating.",
        "Adjustment factors for more than three current-carrying conductors in a raceway or cable are determined from Table 310.15(C)(1).",
        "Per NEC 110.14(C), equipment terminal temperature ratings restrict conductor ampacity; the final allowable ampacity cannot exceed the ampacity for the lowest temperature rating of any connected terminal, device, or conductor.",
    ]


def _build_source_table_from_result(
    result: dict[str, Any],
    calc_mode: str,
):
    size = get_first(result, "selected_size", "conductor_size")
    if not size:
        return None

    return build_nec_table_row_source(
        table_name="table_310_16",
        criteria={"size_awg_kcmil": str(size)},
        columns=[
            "size_awg_kcmil",
            "copper_60c_ampacity",
            "copper_75c_ampacity",
            "copper_90c_ampacity",
            "aluminum_or_copper_clad_aluminum_60c_ampacity",
            "aluminum_or_copper_clad_aluminum_75c_ampacity",
            "aluminum_or_copper_clad_aluminum_90c_ampacity",
        ],
        column_labels={
            "size_awg_kcmil": "Size",
            "copper_60c_ampacity": "Cu 60°C",
            "copper_75c_ampacity": "Cu 75°C",
            "copper_90c_ampacity": "Cu 90°C",
            "aluminum_or_copper_clad_aluminum_60c_ampacity": "Al 60°C",
            "aluminum_or_copper_clad_aluminum_75c_ampacity": "Al 75°C",
            "aluminum_or_copper_clad_aluminum_90c_ampacity": "Al 90°C",
        },
        title="Selected NEC Table 310.16 Ampacity Row",
    )


def _build_report_context(
    result: dict[str, Any],
    calc_mode: str,
) -> dict[str, Any]:
    return {
        "notes": _build_notes(calc_mode),
        "input_pairs": _build_input_pairs(result=result, calc_mode=calc_mode),
        "result_pairs": _build_result_pairs(result),
        "source_table": _build_source_table_from_result(result=result, calc_mode=calc_mode),
    }


def build_word_report(
    result: dict[str, Any],
    calc_mode: str,
) -> bytes:
    doc = init_word_doc(REPORT_TITLE)
    report_context = _build_report_context(result=result, calc_mode=calc_mode)

    add_equations(doc, calc_mode)

    add_bullets(doc, report_context["notes"], heading="Notes and Assumptions")

    add_kv_section_to_word(doc, "Inputs and Parameters Used", report_context["input_pairs"])

    add_source_table_to_word(
        doc,
        report_context["source_table"],
        font_size_header=7,
        font_size=7,
    )

    add_kv_section_to_word(doc, "Results", report_context["result_pairs"])

    return save_word_doc(doc)


def _build_result_summary_pairs(result: dict[str, Any]) -> list[tuple[str, Any]]:
    pairs = [
        ("Final Allowable Ampacity (A)", get_first(result, "calculated_value")),
    ]
    load_current = get_first(result, "load_current")
    if load_current is not None:
        min_size = get_first(result, "min_recommended_size_display", default="—")
        if min_size and min_size != "—":
            pairs.append(("Minimum Required Conductor Size", min_size))
        pairs.append(("Adequate for Load?", yes_no(get_first(result, "is_adequate"))))
    return pairs


def _build_excel_summary_input_pairs(result: dict[str, Any]) -> list[tuple[str, Any]]:
    mat = get_first(result, "material")
    pairs: list[tuple[str, Any]] = [
        ("Conductor Material", MATERIAL_LABELS.get(mat, str(mat))),
    ]
    wire_type = get_first(result, "wire_type")
    if wire_type and wire_type != "Not specified":
        pairs.append(("Wire / Cable Type", wire_type))

    pairs.append(("Selected Conductor Size", get_first(result, "selected_size_display", "selected_size", "conductor_size_display", "conductor_size")))
    pairs.append(("Conductor Insulation Rating (°C)", get_first(result, "temp_rating")))

    terminal_rating = get_first(result, "terminal_temp_rating")
    if terminal_rating is not None and terminal_rating != "—" and terminal_rating != "":
        pairs.append(("Equipment Terminal Rating (°C)", terminal_rating))

    ambient_base = get_first(result, "ambient_base_label") or get_first(result, "ambient_base")
    if ambient_base is not None and ambient_base != "—" and ambient_base != "":
        pairs.append(("Ambient Temperature Base Table", ambient_base))

    pairs.append(("Ambient Operating Temperature (°C)", get_first(result, "ambient_temp_c")))
    pairs.append(("Current-Carrying Conductors", get_first(result, "number_of_conductors")))

    n_parallel = get_first(result, "n_parallel", default=1)
    if n_parallel and n_parallel > 1:
        pairs.append(("Parallel Runs per Phase (N)", n_parallel))

    load_current = get_first(result, "load_current")
    if load_current is not None:
        pairs.append(("Design Load Current (A)", load_current))

    return pairs


def _build_excel_tables_used_pairs(result: dict[str, Any]) -> list[tuple[str, Any]]:
    ambient_base = get_first(result, "ambient_base")
    ambient_table = (
        "Table 310.15(B)(1)(2) — 40°C Base"
        if ambient_base == "40c"
        else "Table 310.15(B)(1)(1) — 30°C Base"
    )
    pairs = [
        ("Ampacity table used", "Table 310.16"),
        ("Ambient correction table used", ambient_table),
    ]
    n_cond = get_first(result, "number_of_conductors")
    if n_cond not in ("1-3", "1", "2", "3"):
        pairs.append(("Conductor adjustment table used", "Table 310.15(C)(1)"))
    if get_first(result, "terminal_temp_rating"):
        pairs.append(("Terminal rating limit basis", "NEC 110.14(C)"))
    return pairs


def _build_excel_correction_factor_rows(result: dict[str, Any]) -> list[tuple[str, Any, str]]:
    ambient_base = get_first(result, "ambient_base")
    ambient_source = "Table 310.15(B)(1)(2)" if ambient_base == "40c" else "Table 310.15(B)(1)(1)"
    n_cond = get_first(result, "number_of_conductors")
    adjustment_source = "Table 310.15(C)(1)" if n_cond not in ("1-3", "1", "2", "3") else "N/A (≤3 conductors)"
    return [
        ("CF_temp (ambient correction)", get_first(result, "ambient_correction"), ambient_source),
        ("AF_cond (conductor adjustment)", get_first(result, "conductor_adjustment"), adjustment_source),
    ]


def _build_excel_ampacity_math_lines(result: dict[str, Any]) -> list[str]:
    table_amp = get_first(result, "table_ampacity")
    cf_temp = get_first(result, "ambient_correction")
    af_cond = get_first(result, "conductor_adjustment")
    derated = get_first(result, "derated_ampacity")
    term_limit = get_first(result, "terminal_limit_ampacity")
    single = get_first(result, "allowable_single")
    n_parallel = get_first(result, "n_parallel", default=1)
    total = get_first(result, "calculated_value")

    lines: list[str] = []
    if None not in (table_amp, cf_temp, af_cond, derated):
        lines.append(
            f"Derated ampacity = I_table x CF_temp x AF_cond = "
            f"{cell_text(table_amp)} x {cell_text(cf_temp)} x {cell_text(af_cond)} = {cell_text(derated)} A"
        )
    if single is not None and derated is not None:
        if term_limit is not None:
            lines.append(
                f"Final allowable ampacity = min(I_derated, I_terminal) = "
                f"min({cell_text(derated)}, {cell_text(term_limit)}) = {cell_text(single)} A"
            )
        else:
            lines.append(f"Final allowable ampacity = I_derated (no terminal limit specified) = {cell_text(single)} A")
    if n_parallel and n_parallel > 1 and single is not None and total is not None:
        lines.append(
            f"Total allowable ampacity = N_parallel x I_allowable = "
            f"{n_parallel} x {cell_text(single)} = {cell_text(total)} A"
        )
    return lines


def build_excel_report(
    result: dict[str, Any],
    calc_mode: str,
) -> bytes:
    wb, ws, row = init_excel_report(REPORT_TITLE, "Summary")

    row = write_kv_sections_to_excel(
        ws,
        row,
        [("Inputs and Parameters Used", _build_excel_summary_input_pairs(result))],
    )

    ws[f"A{row}"] = "Recommended Result"
    ws[f"A{row}"].font = Font(bold=True)
    row += 1
    for label, value in _build_result_summary_pairs(result):
        ws[f"A{row}"] = label
        ws[f"B{row}"] = value
        if label == "Final Allowable Ampacity (A)":
            ws[f"A{row}"].font = Font(bold=True)
            ws[f"B{row}"].font = Font(bold=True)
        row += 1

    autosize_cols(ws)

    ws_tables = wb.create_sheet("Tables Used")
    ws_tables.append(["Item", "Selection"])
    for c in ws_tables[1]:
        c.font = Font(bold=True)
    for label, value in _build_excel_tables_used_pairs(result):
        ws_tables.append([label, value])

    source_table = _build_source_table_from_result(result=result, calc_mode=calc_mode)
    if source_table:
        write_columnar_table(
            ws_tables,
            ws_tables.max_row + 2,
            source_table["title"],
            source_table["rows"],
            source_table["columns"],
        )
    autosize_cols(ws_tables)

    ws_factors = wb.create_sheet("Correction Factors")
    ws_factors.append(["Factor", "Value", "Source"])
    for c in ws_factors[1]:
        c.font = Font(bold=True)
    for label, value, source in _build_excel_correction_factor_rows(result):
        ws_factors.append([label, value, source])
    ws_factors.append([])
    for line in _build_excel_ampacity_math_lines(result):
        ws_factors.append(["Ampacity math", line, "Calculated"])
    autosize_cols(ws_factors)

    return wb_to_bytes(wb)


def render_export_report(
    result: dict[str, Any] | None,
    calc_mode: str,
) -> None:
    can_export = (
        result is not None
        and get_first(result, "calculated_value") is not None
    )

    render_export_buttons(
        prefix="nec_conductors",
        docx_file="nec_conductors_report.docx",
        xlsx_file="nec_conductors_report.xlsx",
        can_export=can_export,
        word_builder=lambda: build_word_report(
            result=result or {},
            calc_mode=calc_mode,
        ),
        excel_builder=lambda: build_excel_report(
            result=result or {},
            calc_mode=calc_mode,
        ),
    )

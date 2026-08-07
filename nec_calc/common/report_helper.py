from __future__ import annotations

import io
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Callable

import streamlit as st
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from lib import nec_tables


# ============================================================
# Value formatting
# ============================================================
def cell_text(val):
    if val is None or val == "":
        return "—"
    if isinstance(val, (int, float)):
        return f"{val:.6g}"
    try:
        return f"{float(val):.6g}"
    except Exception:
        return str(val)


def safe_float_or_text(val):
    if val is None or val == "":
        return None
    try:
        return float(val)
    except Exception:
        return str(val)


def get_first(data, *keys, default=None):
    for key in keys:
        if key in data and data.get(key) is not None:
            return data.get(key)
    return default


def yes_no(value):
    return "Yes" if bool(value) else "No"


def notes_to_pairs(notes):
    return [(f"Note {i + 1}", note) for i, note in enumerate(notes or [])]


# ============================================================
# Word equation helpers
# ============================================================
def omml_r(text: str) -> str:
    return f"<m:r><m:t>{escape(str(text))}</m:t></m:r>"


def omml_sub(base: str, sub: str) -> str:
    return f"<m:sSub><m:e>{omml_r(base)}</m:e><m:sub>{omml_r(sub)}</m:sub></m:sSub>"


def omml_frac(num_inner: str, den_inner: str) -> str:
    return f"<m:f><m:num>{num_inner}</m:num><m:den>{den_inner}</m:den></m:f>"


def omml_sqrt(inner: str) -> str:
    return f'<m:rad><m:degHide m:val="1"/><m:e>{inner}</m:e></m:rad>'


def add_omml_equation_to_paragraph(p, omml_inner: str) -> None:
    xml = f'<m:oMath {nsdecls("m")}>{omml_inner}</m:oMath>'
    p._p.append(parse_xml(xml))


def add_word_equation(doc, label: str, omml_inner: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    add_omml_equation_to_paragraph(p, omml_inner)


# ============================================================
# Word document helpers
# ============================================================
def append_to_value_line(cell, value: str, paragraph_index: int = 1):
    while len(cell.paragraphs) <= paragraph_index:
        cell.add_paragraph("")
    cell.paragraphs[paragraph_index].add_run(value)


def _delete_paragraph(p):
    p._element.getparent().remove(p._element)
    p._p = p._element = None


def remove_leading_blank_paragraphs(doc: Document):
    while doc.paragraphs and doc.paragraphs[0].text.strip() == "":
        _delete_paragraph(doc.paragraphs[0])


def set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    table._tbl.tblPr.append(borders)


def fill_doc_header(doc: Document, title: str) -> None:
    project_number = st.session_state.get("project_number", "")
    designer_name = st.session_state.get("designer_name", "")
    values = {
        (0, 3): project_number,
        (0, 4): "#",
        (2, 3): designer_name,
        (2, 4): datetime.now().strftime("%m/%d/%Y"),
        (3, 2): title,
        (3, 3): "",
        (3, 4): "",
    }
    try:
        hdr_table = doc.sections[0].header.tables[0]
        for (row, col), value in values.items():
            append_to_value_line(hdr_table.cell(row, col), value)
    except Exception:
        doc.add_heading(title, level=0)


def init_word_doc(title: str, template_path: str = "content/files/Template.docx") -> Document:
    path = Path(template_path)
    if not path.exists():
        doc = Document()
        doc.add_heading(title, level=0)
        return doc

    doc = Document(str(path))
    remove_leading_blank_paragraphs(doc)
    fill_doc_header(doc, title)
    return doc


def save_word_doc(doc: Document) -> bytes:
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _set_cell_text(cell, value, font_size, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(cell_text(value))
    run.bold = bold
    run.font.size = Pt(font_size)


def add_word_table(doc, headers, rows, font_size_header=9, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))

    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, font_size_header, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        values = [
            row_data.get(header) if isinstance(row_data, dict) else row_data[i] if i < len(row_data) else None
            for i, header in enumerate(headers)
        ]
        for cell, value in zip(cells, values):
            _set_cell_text(cell, value, font_size)

    set_table_borders(table)
    return table


def add_kv_section_to_word(doc, title, pairs, font_size_header=9, font_size=9):
    if not pairs:
        return None
    doc.add_heading(title, level=1)
    return add_word_table(doc, ["Parameter", "Value"], pairs, font_size_header, font_size)


def add_equations(doc, equations, heading="Equations"):
    if not equations:
        return None
    if heading:
        doc.add_heading(heading, level=1)
    for title, equation in equations:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(equation)


def add_bullets(doc, items, heading="Assumptions", style="CalcBullet"):
    if not items:
        return None
    if heading:
        doc.add_heading(heading, level=1)
    for item in items:
        try:
            doc.add_paragraph(item, style=style)
        except Exception:
            doc.add_paragraph(item)


# ============================================================
# Excel helpers
# ============================================================
def autosize_cols(ws) -> None:
    for col in ws.columns:
        col_letter = get_column_letter(col[0].column)
        max_len = max((len(str(cell.value or "")) for cell in col), default=0)
        ws.column_dimensions[col_letter].width = min(60, max(10, max_len + 2))


def init_excel_report(title: str, sheet_name: str):
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A3"] = "Generated"
    ws["B3"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return wb, ws, 5


def wb_to_bytes(wb: Workbook) -> bytes:
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def write_kv_block(ws, start_row, title, pairs):
    row = start_row
    if not pairs:
        return row
    if title:
        ws[f"A{row}"] = title
        ws[f"A{row}"].font = Font(bold=True)
        row += 1
    for name, val in pairs:
        ws[f"A{row}"] = name
        ws[f"B{row}"] = safe_float_or_text(val)
        row += 1
    return row + 1


def write_kv_sections_to_excel(ws, row, sections):
    for title, pairs in sections:
        row = write_kv_block(ws, row, title, pairs)
    return row


def write_columnar_table(ws, start_row, title, rows, columns):
    row = start_row
    if title:
        ws[f"A{row}"] = title
        ws[f"A{row}"].font = Font(bold=True)
        row += 1

    for col_idx, col_name in enumerate(columns, start=1):
        cell = ws.cell(row=row, column=col_idx)
        cell.value = col_name
        cell.font = Font(bold=True)
    row += 1

    for data_row in rows:
        for col_idx, col_name in enumerate(columns, start=1):
            ws.cell(row=row, column=col_idx).value = safe_float_or_text(data_row.get(col_name))
        row += 1

    return row + 1


def add_table_sheet(wb, sheet_name, columns, rows):
    ws = wb.create_sheet(sheet_name)
    ws.append(list(columns))
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for data_row in rows:
        ws.append([safe_float_or_text(data_row.get(col)) for col in columns])
    autosize_cols(ws)
    return ws


def add_wrapped_list_sheet(wb, sheet_name, header, items, width=110):
    ws = wb.create_sheet(sheet_name)
    ws.append([header])
    ws["A1"].font = Font(bold=True)
    for item in items:
        ws.append([item])
    ws.column_dimensions["A"].width = width
    for row in range(2, 2 + len(items)):
        ws[f"A{row}"].alignment = Alignment(wrap_text=True, vertical="top")
    return ws


def rows_for_excel(table):
    return [{col: row.get(col) for col in table.get("columns", [])} for row in table.get("rows", [])]


# ============================================================
# NEC source table helpers
# ============================================================
def _normalize_nec_table_name(table_name):
    name = str(table_name).strip().upper().replace("-", "_").replace(" ", "_")
    if name.startswith("TABLE_"):
        return name
    return f"TABLE_{name.replace('TABLE', '').replace('_', '').strip()}"


def get_nec_table(table_name):
    """Find a table by name, e.g. "TABLE_310_16" or "table_310_16".

    Goes through the registry rather than the module's attributes so that reports do not
    depend on a table having a module-level constant. Falls back to the flattened registry,
    which is where the sections of a grouped table such as TABLE_4_EMT appear.
    """
    table_id = _normalize_nec_table_name(table_name).lower()
    table = nec_tables.TABLES.get(table_id)
    if table is None:
        table = (nec_tables.get_table_meta(table_id) or {}).get("raw")
    return table


def build_nec_table_row_source(table_name, criteria, columns=None, title=None, column_labels=None):
    if not criteria or any(value is None or value == "" for value in criteria.values()):
        return None

    table = get_nec_table(table_name)
    if table is None:
        return None

    row = nec_tables.get_table_row(table, criteria)
    if row is None:
        return None

    columns = [col for col in (columns or table.get("columns", list(row.keys()))) if col in row]
    if not columns:
        return None

    column_labels = column_labels or {}
    display_columns = [column_labels.get(col, col) for col in columns]
    display_row = {column_labels.get(col, col): row.get(col) for col in columns}

    return {
        "title": title or f"Selected Row — {table.get('title', table_name)}",
        "columns": display_columns,
        "rows": [display_row],
    }


def _source_table_parts(source_table):
    if not source_table:
        return None
    columns = source_table.get("columns", [])
    rows = source_table.get("rows", [])
    if not columns or not rows:
        return None
    return source_table.get("title", "Selected Source Table Row"), columns, rows


def add_source_table_to_word(doc, source_table, font_size_header=7, font_size=7):
    parts = _source_table_parts(source_table)
    if parts is None:
        return None
    title, columns, rows = parts
    doc.add_heading(title, level=1)
    return add_word_table(doc, columns, rows, font_size_header, font_size)


def write_source_table_to_excel(ws, start_row, source_table):
    parts = _source_table_parts(source_table)
    if parts is None:
        return start_row
    title, columns, rows = parts
    return write_columnar_table(ws, start_row, title, rows, columns)


# ============================================================
# Standard report builders
# ============================================================
def can_export_result(result: dict[str, Any] | None, required_keys: tuple[str, ...]) -> bool:
    return result is not None and all(get_first(result, key) is not None for key in required_keys)


def build_standard_word_report(
    report_title: str,
    result: dict[str, Any],
    context_builder: Callable[[dict[str, Any]], dict[str, Any]],
    word_equation_builder: Callable[[Any, dict[str, Any]], None] | None = None,
    notes_heading: str = "Notes and Assumptions",
    input_heading: str = "Inputs and Parameters Used",
    result_heading: str = "Results",
    source_table_font_size_header: int = 7,
    source_table_font_size: int = 7,
) -> bytes:
    doc = init_word_doc(report_title)
    context = context_builder(result)

    if word_equation_builder is not None:
        word_equation_builder(doc, context)
    elif context.get("equations"):
        add_equations(doc, context["equations"], heading="Equations Used")

    add_bullets(doc, context.get("notes"), heading=notes_heading)
    add_kv_section_to_word(doc, input_heading, context.get("input_pairs"))
    add_source_table_to_word(
        doc,
        context.get("source_table"),
        font_size_header=source_table_font_size_header,
        font_size=source_table_font_size,
    )
    add_kv_section_to_word(doc, result_heading, context.get("result_pairs"))

    return save_word_doc(doc)


def build_standard_excel_report(
    report_title: str,
    sheet_name: str,
    result: dict[str, Any],
    context_builder: Callable[[dict[str, Any]], dict[str, Any]],
    excel_equation_builder: Callable[[dict[str, Any]], list[tuple[str, Any]]] | None = None,
    notes_heading: str = "Notes and Assumptions",
    input_heading: str = "Inputs and Parameters Used",
    result_heading: str = "Results",
) -> bytes:
    context = context_builder(result)
    wb, ws, row = init_excel_report(report_title, sheet_name)

    equations = excel_equation_builder(context) if excel_equation_builder is not None else context.get("equations", [])

    row = write_kv_sections_to_excel(
        ws,
        row,
        [
            ("Equations Used", equations),
            (notes_heading, notes_to_pairs(context.get("notes", []))),
            (input_heading, context.get("input_pairs", [])),
        ],
    )
    row = write_source_table_to_excel(ws, row, context.get("source_table"))
    write_kv_sections_to_excel(ws, row, [(result_heading, context.get("result_pairs", []))])

    autosize_cols(ws)
    return wb_to_bytes(wb)


def render_standard_export_report(
    prefix: str,
    docx_file: str,
    xlsx_file: str,
    result: dict[str, Any] | None,
    required_keys: tuple[str, ...],
    word_builder: Callable[[dict[str, Any]], bytes],
    excel_builder: Callable[[dict[str, Any]], bytes],
) -> None:
    render_export_buttons(
        prefix=prefix,
        docx_file=docx_file,
        xlsx_file=xlsx_file,
        can_export=can_export_result(result, required_keys),
        word_builder=lambda: word_builder(result or {}),
        excel_builder=lambda: excel_builder(result or {}),
    )


# ============================================================
# Export button pair
# ============================================================
def _render_export_button(
    prefix,
    ext,
    label,
    file_name,
    mime,
    builder,
    can_export,
):
    build_key = f"{prefix}_build_{ext}"
    bytes_key = f"{prefix}_{ext}_bytes"
    download_key = f"{prefix}_download_{ext}"

    if st.button(f"Prepare {label} report (.{ext})", key=build_key):
        try:
            st.session_state[bytes_key] = builder()
            st.success(f"{label} report prepared. Use the download button below.")
        except Exception as e:
            st.error(f"Failed to build {label} report: {e}")

    file_bytes = st.session_state.get(bytes_key)

    st.download_button(
        f"⬇️ Download {label} report (.{ext})",
        data=file_bytes or b"",
        file_name=file_name,
        mime=mime,
        disabled=(not can_export) or (file_bytes is None),
        key=download_key,
    )


def render_export_buttons(prefix, docx_file, xlsx_file, can_export, word_builder, excel_builder):
    exports = [
        (
            "docx",
            "Word",
            docx_file,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            word_builder,
        ),
        (
            "xlsx",
            "Excel",
            xlsx_file,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            excel_builder,
        ),
    ]

    for col, export in zip(st.columns([1, 1], gap="large"), exports):
        with col:
            _render_export_button(
                prefix,
                *export,
                can_export=can_export,
            )


__all__ = [
    # Value formatting
    "cell_text",
    "safe_float_or_text",
    "get_first",
    "yes_no",
    "notes_to_pairs",

    # Word equation helpers
    "omml_r",
    "omml_sub",
    "omml_frac",
    "omml_sqrt",
    "add_omml_equation_to_paragraph",
    "add_word_equation",

    # Word document helpers
    "append_to_value_line",
    "remove_leading_blank_paragraphs",
    "set_table_borders",
    "fill_doc_header",
    "init_word_doc",
    "save_word_doc",
    "add_word_table",
    "add_kv_section_to_word",
    "add_equations",
    "add_bullets",

    # Excel helpers
    "autosize_cols",
    "init_excel_report",
    "wb_to_bytes",
    "write_kv_block",
    "write_kv_sections_to_excel",
    "write_columnar_table",
    "add_table_sheet",
    "add_wrapped_list_sheet",
    "rows_for_excel",

    # NEC source table helpers
    "get_nec_table",
    "build_nec_table_row_source",
    "add_source_table_to_word",
    "write_source_table_to_excel",

    # Standard report builders
    "can_export_result",
    "build_standard_word_report",
    "build_standard_excel_report",
    "render_standard_export_report",

    # Export button pair
    "render_export_buttons",
]
